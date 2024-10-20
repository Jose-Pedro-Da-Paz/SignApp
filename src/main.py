import os
from tkinter import Tk, Label, Entry, Button, filedialog, messagebox, PhotoImage, font, Toplevel, Scale
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.shared import Inches  # Para ajustar o tamanho da imagem
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email import encoders
from assinatura import SignatureApp

def preencher_certificado(caminho_documento, caminho_saida, dados):
    try:
        documento = Document(caminho_documento)

        # Substituição dos campos no documento
        preenchimentos = {
            '[NOME]': dados['nome'],
            '[CIM]': dados['cim'],
            '[LOJA]': dados['loja'],
            '[N_LOJA]': dados['n_loja'],
            '[POTENCIA]': dados['potencia'],
            '[SESSAO]': dados['sessao'],
            '[DATA]': dados['data']
        }

        for paragrafo in documento.paragraphs:
            for chave, valor in preenchimentos.items():
                if chave in paragrafo.text:
                    paragrafo.text = paragrafo.text.replace(chave, valor)

            # Padronizar a fonte para Times New Roman, tamanho 14
            for run in paragrafo.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(14)
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

        # Inserir as assinaturas no documento
        for paragrafo in documento.paragraphs:
            if '[ASS1]' in paragrafo.text:
                paragrafo.clear()  # Remove o texto
                run = paragrafo.add_run()
                run.add_picture(dados['ass1'], width=Inches(2))  # Ajusta o tamanho da imagem

        documento.save(caminho_saida)
        messagebox.showinfo("Sucesso", f"Documento salvo em {caminho_saida}")
    except Exception as e:
        messagebox.showerror("Erro", f"Ocorreu um erro: {e}")

def enviar_email_com_anexo(caminho_arquivo, destinatario):
    remetente = '***************'
    assunto = 'Certificado de Visitação'
    corpo_email = 'Segue em anexo o certificado de visitação.'

    mensagem = MIMEMultipart()
    mensagem['From'] = remetente
    mensagem['To'] = destinatario
    mensagem['Subject'] = assunto
    mensagem.attach(MIMEText(corpo_email, 'plain'))

    with open(caminho_arquivo, 'rb') as anexo:
        mime_base = MIMEBase('application', 'octet-stream')
        mime_base.set_payload(anexo.read())
        encoders.encode_base64(mime_base)
        mime_base.add_header('Content-Disposition', f'attachment; filename={os.path.basename(caminho_arquivo)}')
        mensagem.attach(mime_base)

    servidor_smtp = 'smtp.gmail.com'
    porta = 587
    usuario = remetente
    senha = '************'  # Substitua por sua senha de app

    try:
        servidor = smtplib.SMTP(servidor_smtp, porta)
        servidor.starttls()
        servidor.login(usuario, senha)
        servidor.send_message(mensagem)
        servidor.quit()
        messagebox.showinfo("Sucesso", "E-mail enviado com sucesso!")
    except Exception as e:
        messagebox.showerror("Erro", f"Erro ao enviar e-mail: {e}")

def gerar_documento():
    nome = nome_entry.get()
    cim = cim_entry.get()
    loja = loja_entry.get()
    n_loja = n_loja_entry.get()
    potencia = potencia_entry.get()
    sessao = sessao_entry.get()
    data = data_entry.get()
    ass1 = ass1_entry.get()

    if not nome or not cim or not loja or not n_loja or not potencia or not sessao or not data or not ass1:
        messagebox.showerror("Erro", "Todos os campos são obrigatórios.")
        return

    caminho_documento = 'G:\\IAs_e _MachineLearning\\SignLoja\\CERTIFICADO_VISITACAO.docx'
    caminho_saida = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Documentos do Word", "*.docx")])
    if caminho_documento and caminho_saida:
        preencher_certificado(caminho_documento, caminho_saida, {
            'nome': nome, 'cim': cim, 'loja': loja, 'n_loja': n_loja,
            'potencia': potencia, 'sessao': sessao, 'data': data, 'ass1': ass1
        })

def enviar_documento():
    destinatario = email_entry.get()
    caminho_arquivo = filedialog.askopenfilename(filetypes=[("Documentos do Word", "*.docx")])
    if caminho_arquivo and destinatario:
        enviar_email_com_anexo(caminho_arquivo, destinatario)
    else:
        messagebox.showerror("Erro", "Por favor, selecione um arquivo e insira um e-mail.")

def ajustar_tamanho_fonte(novo_tamanho):
    fonte_nova = font.Font(size=int(novo_tamanho))
    nome_entry.configure(font=fonte_nova)
    cim_entry.configure(font=fonte_nova)
    loja_entry.configure(font=fonte_nova)
    n_loja_entry.configure(font=fonte_nova)
    potencia_entry.configure(font=fonte_nova)
    sessao_entry.configure(font=fonte_nova)
    data_entry.configure(font=fonte_nova)
    ass1_entry.configure(font=fonte_nova)
    email_entry.configure(font=fonte_nova)

def abrir_configuracoes():
    # Criar uma nova janela
    config_window = Toplevel(app)
    config_window.title("Configurações de Fonte")
    config_window.geometry("300x200")

    # Rótulo explicando a funcionalidade
    Label(config_window, text="Selecione o tamanho da fonte:").pack(pady=10)

    # Scale para ajustar o tamanho da fonte
    scale = Scale(config_window, from_=8, to=48, orient='horizontal', command=ajustar_tamanho_fonte)
    scale.set(14)  # Tamanho inicial da fonte
    scale.pack(pady=20)

    # Botão para aplicar as configurações
    Button(config_window, text="Aplicar", command=lambda: ajustar_tamanho_fonte(scale.get())).pack(pady=10)

def abrir_signature_app():
    # Comando para executar o script de captura de assinatura
    root = Tk()
    sign = SignatureApp(root)
    root.mainloop()
    
# Função para selecionar imagem de assinatura
def selecionar_imagem(entry):
    caminho_imagem = filedialog.askopenfilename(filetypes=[("Imagens", "*.png;*.jpg;*.jpeg")])
    if caminho_imagem:
        entry.delete(0, 'end')  # Limpa o campo
        entry.insert(0, caminho_imagem)  # Insere o caminho da assinatura no campo correspondente

app = Tk()
app.title("Gerador de Certificado")

# Definindo ícone da janela
app.iconphoto(False, PhotoImage(file='G:\\IAs_e _MachineLearning\\SignLoja\\logotipo.png'))

# Carregar a imagem da logo
logo_path = "G:\\IAs_e _MachineLearning\\SignLoja\\logotipo.png"  # Certifique-se de que sua imagem está no formato GIF
logo_photo = PhotoImage(file=logo_path)
logo_photo = logo_photo.subsample(2, 2)

# Adicionar a logo na interface 
logo_label = Label(app, image=logo_photo)
logo_label.grid(row=0, column=0, columnspan=2, pady=10)

# Campos de entrada
Label(app, text="Nome Ir∴").grid(row=1, column=0, sticky='w')
nome_entry = Entry(app)
nome_entry.grid(row=1, column=1, sticky='we')

Label(app, text="CIM").grid(row=2, column=0, sticky='w')
cim_entry = Entry(app)
cim_entry.grid(row=2, column=1, sticky='we')

Label(app, text="Nome da Loja").grid(row=3, column=0, sticky='w')
loja_entry = Entry(app)
loja_entry.grid(row=3, column=1, sticky='we')

Label(app, text="Número da Loja").grid(row=4, column=0, sticky='w')
n_loja_entry = Entry(app)
n_loja_entry.grid(row=4, column=1, sticky='we')

Label(app, text="Potência").grid(row=5, column=0, sticky='w')
potencia_entry = Entry(app)
potencia_entry.grid(row=5, column=1, sticky='we')

Label(app, text="Sessão").grid(row=6, column=0, sticky='w')
sessao_entry = Entry(app)
sessao_entry.grid(row=6, column=1, sticky='we')

Label(app, text="Data").grid(row=7, column=0, sticky='w')
data_entry = Entry(app)
data_entry.grid(row=7, column=1, sticky='we')

Label(app, text="Assinatura 1").grid(row=8, column=0, sticky='w')
ass1_entry = Entry(app)
ass1_entry.grid(row=8, column=1, sticky='we')
Button(app, text="Selecionar Imagem", command=lambda: selecionar_imagem(ass1_entry)).grid(row=8, column=2)

Label(app, text="E-mail").grid(row=10, column=0, sticky='w')
email_entry = Entry(app)
email_entry.grid(row=10, column=1, sticky='we')

# Botões de ação
Button(app, text="Gerar Certificado", command=gerar_documento).grid(row=11, column=0, pady=10)
Button(app, text="Enviar por E-mail", command=enviar_documento).grid(row=11, column=1, pady=10)

# Botão para abrir configurações
Button(app, text="Configurações", command=abrir_configuracoes).grid(row=12, column=0, pady=10)

# Botão para capturar assinatura
Button(app, text="Capturar Assinatura", command=abrir_signature_app).grid(row=12, column=1, pady=10)

# Expande os campos de entrada
app.grid_columnconfigure(1, weight=1)

app.mainloop()
